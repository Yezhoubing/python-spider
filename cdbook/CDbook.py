# -*- coding:utf-8 -*-
# @author:Ye Zhoubing
# @datetime:2023/12/9 17:45
# @software: PyCharm

"""
CDBOOK爬取小说（本地）
"""
# -*- codeing=utf-8 -*-
# @Time:2021/5/27 20:17
# @Author:Ye Zhoubing
# @File: novel_new.py
# @software:PyCharm
# 内容在span标签内

from bs4 import BeautifulSoup
from docx import Document   # 安装docx的命令是：pip install python-docx
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import wordTools


# 打开本地html文件
with open(r'C:\Users\Lenovo\Desktop\git\cdbook\1.html', encoding='utf-8') as f:
    html = f.read()

soup = BeautifulSoup(html, 'html.parser')

meta_tag = soup.find('span', id="thread_subject").text  # 空格问题没有解决

document = Document()

# 有的文章是style='font-size:14px'、有的是color="#000"，todo:有的文章这三种都不是
article = soup.find_all('font', style='font-size:14px')
if(len(article) == 0):
    article = soup.find_all('font', color="#000")
else:
    pass
[s.extract() for s in soup.find_all('font',{'class':'jammer'})]  # 去掉乱码内容
[s.extract() for s in soup.find_all('span',style="display:none")]  # 去掉乱码内容
for item in article:
    # regStr = ".*?([\u4E00-\u9FA5]+).*?" #同时会把br，标点符号去掉
    # data = re.findall(regStr, item.text)
    # #re.sub()表示替换,\u4E00-\u9FA5是中文编码范围
    print(item.text)

    # 设置字体样式及大小，需要先设置好字体，再添加内容，否则设置的就是无效的.todo:待运行
    """
    设置正文字型
        英文字型:Times New Roman
        中文字型:宋体
    """
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'  # 必须先设置font.name
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    """设置正文字体的大小"""
    style.font.size = Pt(10.5)  # 10.5磅字，对应五号字
    # """设置正文字体颜色"""
    # style.font.color.rgb = RGBColor(0, 0, 0)

    # 为文档添加一个段落
    paragraph = document.add_paragraph()

    # 添加一个文本
    run = paragraph.add_run(item.text)


document.save(f'{meta_tag}.docx')

# 清除word文档中的多余硬回车（就是向下的箭头）
word_file = r'C:\Users\Lenovo\Desktop\git\cdbook\{}.docx'.format(meta_tag)
paragraphs = wordTools.read_word_document(word_file)
cleaned_paragraphs = wordTools.remove_extra_newlines(paragraphs)
wordTools.save_modified_document(cleaned_paragraphs, word_file)

print(meta_tag + "over")