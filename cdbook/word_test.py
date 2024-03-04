# -*- coding:utf-8 -*-
# @author:Ye Zhoubing
# @datetime:2024/3/4 18:52
# @software: PyCharm
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

document = Document()

# # 直接设置中文字体，对中文无效
# paragraph1 = document.add_paragraph()
# run = paragraph1.add_run('aBCDefg这是中文')
# font = run.font
# font.name = '宋体'

# 方法1 直接修改一个已有样式的所有文字的样式
style = document.styles['Normal']
style.font.name = 'Times New Roman' # 必须先设置font.name
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

paragraph1 = document.add_paragraph()
run = paragraph1.add_run('修改Normal，修改所有字体')

document.save('test.docx')