# -*- codeing=utf-8 -*-
# @Time:2021/5/27 20:17
# @Author:Ye Zhoubing
# @File: novel_new.py
# @software:PyCharm
# 注意：这个程序生成了文件后如果再次运行，会报错，就是同名文件的原因
""""
    添加的图片有三种情况：
    1.封面图
    2.在文本div中的图
    3.在一个单独的div中的图
"""
#从本地文件中提取小说
import os
from bs4 import BeautifulSoup
from docx import Document
import re
import shutil
import traceback
import requests
from docx.shared import Inches


def main():
    path = os.path.abspath("提取小说.py")  # 得到这个文件路径
    path = os.path.dirname(path)  # 提取最后一个/之前的内容

    #确定文件夹下文件数
    dir = r"C:\Users\32649\Desktop\novel"
    # print(os.listdir(dir))

    # print(path)
    # 准备文档
    document = Document()
    # 在Python中\是转义符，\u表示其后是UNICODE编码，因此\User在这里会报错，在字符串前面加个r表示就可以了
    for num in os.listdir(dir):
        local_file = dir+r'\{}'.format(num)
        with open(local_file,encoding="utf-8") as f:
            content=f.read()
        soup = BeautifulSoup(content,"lxml")
        [s.extract() for s in soup('span')] #去掉span标签及内容
        # font需要删除class="jammer"的乱码标签
        [s.extract() for s in soup.find_all('font',{'class':'jammer'})] #去掉span标签及内容


        article=soup.find_all("div",class_="t_f")
        # document.add_picture("https://www.jingjiniao.info/data/attachment/forum/202103/18/134424hbrrlmrqlal42kzx.png")

        #添加封面图
        thumb_img = soup.find("img", onload="thumbImg(this)")
        if (thumb_img):
            single_url = "https://www.jingjiniao.info/" + thumb_img.get("src")  # 得到封面图
            # 保存封面图片为thumb.jpg文件
            with open("thumb.jpg", 'wb') as f:
                response = requests.get(single_url).content
                f.write(response)
            # 将图片加入word文档中
            document.add_picture("thumb.jpg", width=Inches(6))  # 定义一下大小，建议只设置宽度，会等比例压缩；两者都设置可能图片大小就很怪

            # 删除下载的图片
            os.remove("thumb.jpg")

        name = soup.find("title").text
        obj = re.compile(r"- (?P<f_name>.*?)-",re.S)
        result = obj.search(name)
        f_name = result.group('f_name')
        f_name = f_name.replace(u'\u3000','')
        f_name = f_name.replace('/','')
        print(f_name)


        for item in article:
            # regStr = ".*?([\u4E00-\u9FA5]+).*?" #同时会把br，标点符号去掉
            # data = re.findall(regStr, item.text)
            # #re.sub()表示替换,\u4E00-\u9FA5是中文编码范围
            print(item.text)
            document.add_paragraph(item.text)

            # 往word中添加图片，这个针对的是在文本中的图片
            img_url = item.find_all("img")
            if (img_url):
                # 文章源文件html的图片保存在zoomfile中,这里将所有图片保存在对应的文后
                for single_url in img_url:
                    # img_url = "https://www.jingjiniao.info/" + img_url.get("src")  # 得到图
                    # 表情包也会被识别为图片,此时的链接属性为src，不是zoomfile
                    status = single_url.get("zoomfile")
                    if(status):
                        single_url = "https://www.jingjiniao.info/" + single_url.get("zoomfile")  # 得到图
                        # 保存图片为picture.jpg文件
                        with open("picture.jpg", 'wb') as f:
                            response = requests.get(single_url).content
                            f.write(response)
                            # 将图片加入word文档中
                        document.add_picture("picture.jpg", width=Inches(6))  # 定义一下大小，建议只设置宽度，会等比例压缩；两者都设置可能图片大小就很怪
                        # 删除下载的图片
                        os.remove("picture.jpg")

                    else:
                        # single_url = "https://www.jingjiniao.info/" + single_url.get("src")
                        # # 保存图片为picture.jpg文件
                        # with open("picture.jpg", 'wb') as f:
                        #     response = requests.get(single_url).content
                        #     f.write(response)
                        # # 将图片加入word文档中
                        # document.add_picture("picture.jpg", width=Inches(6))  # 定义一下大小，建议只设置宽度，会等比例压缩；两者都设置可能图片大小就很怪
                        pass  # TODO word中保存字符串以后解决




        # 往word添加单独div中的图片,这样只能单独添加在末尾
        div_imgs = soup.find_all("div",class_="mbn savephotop")
        if(div_imgs):
            for div_img in div_imgs:
                div_img = div_img.find("img").get("zoomfile")
                div_url = "https://www.jingjiniao.info/" + div_img
                # 保存封面图片为picture.jpg文件
                with open("div.jpg", 'wb') as f:
                    response = requests.get(div_url).content
                    f.write(response)
                # 将图片加入word文档中
                document.add_picture("div.jpg", width=Inches(6))  # 定义一下大小，建议只设置宽度，会等比例压缩；两者都设置可能图片大小就很怪

                # 删除下载的图片
                os.remove("div.jpg")

        # 保存.docx文档
        # document.save(f'F:\视频\{f_name}.docx') #保存绝对路径若为c盘，需要管理员权限，可以选择保存在其它位置
    document.save(f'{f_name}.docx') #保存在相对路径
    move_file(r'C:\Users\32649\Desktop\pythonProject',r'C:\Users\32649\Desktop\pythonProject\荆棘鸟翻译',f'{f_name}.docx')
    return f_name

def move_file(src_path, dst_path, file):
    #移动文件
    print('from : ', src_path)
    print('to : ', dst_path)
    try:
        # cmd = 'chmod -R +x ' + src_path
        # os.popen(cmd)
        f_src = os.path.join(src_path, file)
        if not os.path.exists(dst_path):
            os.mkdir(dst_path)
        f_dst = os.path.join(dst_path, file)
        # 判断是否有同名文件，有则新的文件
        if os.path.exists(f_dst):
            os.remove(f_dst)
            shutil.move(f_src, f_dst)
        else:
            shutil.move(f_src, f_dst)
    except Exception as e:
        print('move_file ERROR: ', e)
        traceback.print_exc()

if __name__ == '__main__':
    file_name = main()
    print(file_name,'完成')

