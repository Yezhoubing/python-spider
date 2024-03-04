# -*- coding:utf-8 -*-
# @author:Ye Zhoubing
# @datetime:2023/12/17 15:24
# @software: PyCharm
import docx

def read_word_document(file_path):
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return paragraphs

def remove_extra_newlines(paragraphs):
    cleaned_paragraphs = []
    for paragraph in paragraphs:
        cleaned_paragraph = paragraph.replace("\n", " ")
        cleaned_paragraphs.append(cleaned_paragraph)
    return cleaned_paragraphs

def save_modified_document(paragraphs, output_file_path):
    doc = docx.Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(output_file_path)