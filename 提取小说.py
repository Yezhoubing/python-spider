# -*- codeing=utf-8 -*-
# @Time:2021/3/24 14:21
# @Author:Ye Zhoubing
# @File: 提取小说.py
# @software:PyCharm

#从本地文件中提取小说
import os
from bs4 import BeautifulSoup
import urllib.request
from docx import Document

def main():
    path = os.path.abspath("提取小说.py")  # 得到这个文件路径
    path = os.path.dirname(path)  # 提取最后一个/之前的内容
    # print(path)
    # 在Python中\是转义符，\u表示其后是UNICODE编码，因此\User在这里会报错，在字符串前面加个r表示就可以了
    with open(r"C:\Users\yezhoubing\Desktop\小说\novel.html",encoding="utf-8") as f:
        content=f.read()
    soup=BeautifulSoup(content,"lxml")
    article=soup.find_all("div",align="left")
    document = Document()
    # document.add_picture("https://www.jingjiniao.info/data/attachment/forum/202103/18/134424hbrrlmrqlal42kzx.png")
    for item in article:

        print(item.text)


        #保存为word文档



        p=document.add_paragraph(item.text)


    # 保存.docx文档
    document.save(f'{path}\荆棘鸟.docx')








if __name__ == '__main__':
    main()

